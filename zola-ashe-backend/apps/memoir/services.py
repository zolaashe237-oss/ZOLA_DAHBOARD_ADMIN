"""Génération du document Word pour un mémoire autobiographique soumis."""
from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

from django.utils import timezone

if TYPE_CHECKING:
    from .models import MemoirDraft

from .questions import PARTS


def _has_content(answer: dict) -> bool:
    if not isinstance(answer, dict):
        return False
    structured = answer.get("structured") or {}
    return bool(
        (answer.get("text") or "").strip()
        or (answer.get("audioTranscript") or "").strip()
        or answer.get("notApplicable")
        or any(
            (isinstance(v, list) and v) or (isinstance(v, str) and v.strip())
            for v in structured.values()
        )
    )


def generate_memoir_docx(draft: "MemoirDraft") -> bytes:
    """Construit un fichier .docx à partir des réponses du brouillon soumis."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Marges ──────────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Styles de base ──────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Garamond"
    normal.font.size = Pt(11)

    def _gold():
        return RGBColor(0xC9, 0xA2, 0x27)

    def _dark():
        return RGBColor(0x1A, 0x1A, 0x2E)

    # ── Page de couverture ──────────────────────────────────────────────────────
    user = draft.user
    member_name = getattr(user, "full_name", None) or user.email

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("ÉDITIONS ZOLA ASHÉ")
    run.font.size = Pt(10)
    run.font.color.rgb = _gold()
    run.font.bold = True

    doc.add_paragraph()

    title_para2 = doc.add_paragraph()
    title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title_para2.add_run("Mon Histoire")
    run2.font.size = Pt(28)
    run2.font.bold = True
    run2.font.color.rgb = _dark()

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = sub_para.add_run("Récit de vie autobiographique")
    run3.font.size = Pt(13)
    run3.font.italic = True
    run3.font.color.rgb = _gold()

    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(member_name).font.size = Pt(14)

    meta_lines = []
    if getattr(user, "email", None):
        meta_lines.append(user.email)
    if getattr(user, "phone", None):
        meta_lines.append(user.phone)
    if getattr(user, "country", None):
        meta_lines.append(user.country)

    if meta_lines:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_meta = meta_para.add_run(" · ".join(meta_lines))
        run_meta.font.size = Pt(10)
        run_meta.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    submitted_at = draft.submitted_at or timezone.now()
    date_str = submitted_at.strftime("%d %B %Y")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = date_para.add_run(f"Soumis le {date_str}")
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ── Contenu par chapitre ────────────────────────────────────────────────────
    answers: dict = draft.answers if isinstance(draft.answers, dict) else {}

    for part in PARTS:
        # Collect answered questions for this part
        answered = [
            (qid, qtext, answers[qid])
            for qid, qtext in part["questions"]
            if qid in answers and _has_content(answers[qid])
        ]
        if not answered:
            continue

        # Chapter heading
        heading = doc.add_heading(level=1)
        heading.clear()
        run_h = heading.add_run(f"Chapitre {part['id']}  —  {part['title']}")
        run_h.font.color.rgb = _gold()
        run_h.font.size = Pt(16)
        run_h.font.bold = True

        for qid, qtext, answer in answered:
            # Question heading
            q_heading = doc.add_heading(level=2)
            q_heading.clear()
            run_q = q_heading.add_run(qtext)
            run_q.font.size = Pt(11)
            run_q.font.bold = True
            run_q.font.color.rgb = _dark()

            # Non concerné
            if answer.get("notApplicable"):
                p = doc.add_paragraph()
                run_na = p.add_run("(Non concerné(e))")
                run_na.font.italic = True
                run_na.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                doc.add_paragraph()
                continue

            # Champs structurés
            structured = answer.get("structured") or {}
            filled = {
                k: v for k, v in structured.items()
                if (isinstance(v, list) and v) or (isinstance(v, str) and v.strip())
            }
            if filled:
                table = doc.add_table(rows=len(filled), cols=2)
                table.style = "Table Grid"
                for i, (k, v) in enumerate(filled.items()):
                    row = table.rows[i]
                    label_cell = row.cells[0]
                    label_cell.width = Cm(5)
                    label_run = label_cell.paragraphs[0].add_run(k.replace("_", " ").capitalize())
                    label_run.font.bold = True
                    label_run.font.size = Pt(10)
                    label_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                    val_cell = row.cells[1]
                    val_text = ", ".join(v) if isinstance(v, list) else str(v)
                    val_run = val_cell.paragraphs[0].add_run(val_text)
                    val_run.font.size = Pt(10)
                doc.add_paragraph()

            # Texte libre
            text = (answer.get("text") or "").strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)

            # Transcription audio
            transcript = (answer.get("audioTranscript") or "").strip()
            if transcript:
                label_p = doc.add_paragraph()
                run_label = label_p.add_run("Témoignage oral (transcription) :")
                run_label.font.bold = True
                run_label.font.size = Pt(9)
                run_label.font.color.rgb = RGBColor(0x44, 0x88, 0xCC)

                t_p = doc.add_paragraph(transcript)
                t_p.paragraph_format.left_indent = Cm(1)
                run_t = t_p.runs[0] if t_p.runs else t_p.add_run(transcript)
                run_t.font.italic = True
                run_t.font.size = Pt(10)

            doc.add_paragraph()

        doc.add_page_break()

    # ── Pied de document ────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer_para.add_run("— Éditions ZOLA ASHÉ · Collection Mémoires —")
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = _gold()
    run_footer.font.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
